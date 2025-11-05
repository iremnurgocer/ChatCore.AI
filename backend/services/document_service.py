# -*- coding: utf-8 -*-
"""
Module: Document Service
Description: File parsing (PDF, DOCX, XLSX, TXT), text extraction, chunking, and FAISS indexing.
"""
import os
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import asyncio

# Document parsing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# LangChain imports
try:
    from langchain_core.documents import Document as LangChainDocument
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

from core.config import get_settings
from core.logger import APILogger, ErrorCategory
from models.document_model import Document

settings = get_settings()

# Upload directory
UPLOAD_DIR = Path(__file__).parent.parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Max file size: 50MB
MAX_FILE_SIZE = 50 * 1024 * 1024

# Allowed file types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt"}


class DocumentService:
    """Service for document parsing, chunking, and indexing"""
    
    def __init__(self):
        self.text_splitter = None
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
    
    async def parse_file(
        self,
        file_path: Path,
        file_type: str
    ) -> Tuple[str, Dict]:
        """
        Parse file and extract text
        
        Returns:
            (text_content, metadata)
        """
        if file_type == "pdf":
            return await self._parse_pdf(file_path)
        elif file_type == "docx":
            return await self._parse_docx(file_path)
        elif file_type == "xlsx":
            return await self._parse_xlsx(file_path)
        elif file_type == "txt":
            return await self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _parse_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Parse PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        def _extract_text():
            text_parts = []
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get("/Title", "") if pdf_reader.metadata else ""
                }
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"Page {page_num + 1}:\n{text}")
                    except Exception as e:
                        APILogger.log_error(
                            "/document/parse",
                            e,
                            None,
                            ErrorCategory.AI_ERROR,
                            page=page_num
                        )
            
            return "\n\n".join(text_parts), metadata
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        text, metadata = await loop.run_in_executor(None, _extract_text)
        
        return text, metadata
    
    async def _parse_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Parse DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        def _extract_text():
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "title": doc.core_properties.title or ""
            }
            
            return text, metadata
        
        loop = asyncio.get_event_loop()
        text, metadata = await loop.run_in_executor(None, _extract_text)
        
        return text, metadata
    
    async def _parse_xlsx(self, file_path: Path) -> Tuple[str, Dict]:
        """Parse XLSX file"""
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
        
        def _extract_text():
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"Sheet: {sheet_name}"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                text_parts.append("\n".join(sheet_text))
            
            metadata = {
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames
            }
            
            return "\n\n".join(text_parts), metadata
        
        loop = asyncio.get_event_loop()
        text, metadata = await loop.run_in_executor(None, _extract_text)
        
        return text, metadata
    
    async def _parse_txt(self, file_path: Path) -> Tuple[str, Dict]:
        """Parse TXT file"""
        def _extract_text():
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            
            metadata = {
                "line_count": len(text.splitlines()),
                "char_count": len(text)
            }
            
            return text, metadata
        
        loop = asyncio.get_event_loop()
        text, metadata = await loop.run_in_executor(None, _extract_text)
        
        return text, metadata
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[LangChainDocument]:
        """
        Split text into chunks for indexing
        
        Returns:
            List of LangChain Document objects
        """
        if not LANGCHAIN_AVAILABLE:
            # Fallback: simple chunking
            chunks = []
            words = text.split()
            current_chunk = []
            current_size = 0
            
            for word in words:
                word_size = len(word) + 1  # +1 for space
                if current_size + word_size > chunk_size and current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = current_chunk[-chunk_overlap:] if chunk_overlap > 0 else []
                    current_size = sum(len(w) + 1 for w in current_chunk)
                
                current_chunk.append(word)
                current_size += word_size
            
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            
            return [
                LangChainDocument(page_content=chunk, metadata={})
                for chunk in chunks
            ]
        
        # Use LangChain text splitter
        if not self.text_splitter:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len
            )
        
        chunks = self.text_splitter.split_text(text)
        
        return [
            LangChainDocument(page_content=chunk, metadata={})
            for chunk in chunks
        ]
    
    def validate_file(
        self,
        file_name: str,
        file_size: int
    ) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file
        
        Returns:
            (is_valid, error_message)
        """
        # Check file size
        if file_size > MAX_FILE_SIZE:
            return False, f"File size exceeds maximum ({MAX_FILE_SIZE / 1024 / 1024}MB)"
        
        # Check extension
        ext = Path(file_name).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return False, f"File type not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        
        return True, None
    
    def get_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def save_uploaded_file(
        self,
        file_content: bytes,
        file_name: str,
        user_id: int
    ) -> Path:
        """
        Save uploaded file to disk
        
        Returns:
            Path to saved file
        """
        # Create user-specific directory
        user_dir = UPLOAD_DIR / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in file_name if c.isalnum() or c in "._-")
        unique_name = f"{timestamp}_{safe_name}"
        
        file_path = user_dir / unique_name
        
        # Save file
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        return file_path


# Global instance
document_service = DocumentService()

